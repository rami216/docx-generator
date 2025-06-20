from flask import Flask, request, send_file
from docx import Document
from docx.shared import Pt, Inches
import os
import json
import logging # Added back logging for better debugging

app = Flask(__name__)
logging.basicConfig(level=logging.INFO) # Configure basic logging

@app.route("/generate-docx", methods=["POST"])
def generate_docx():
    app.logger.info("--- Incoming Request ---") # Added logging
    app.logger.info(f"Headers: {request.headers}") # Added logging
    app.logger.info(f"Content-Type: {request.headers.get('Content-Type')}") # Added logging

    # Safely parse JSON whether it's a dict or string
    data = request.get_json()
    if isinstance(data, str):
        try:
            data = json.loads(data)
        except json.JSONDecodeError as e:
            app.logger.error(f"Error decoding top-level JSON: {e}")
            app.logger.error(f"Raw data received: {data}")
            return {"error": f"Invalid top-level JSON payload from n8n: {e}"}, 400
    
    # Ensure data is a dictionary after parsing
    if not isinstance(data, dict):
        app.logger.error(f"Final data is not a dictionary: {type(data)}. Value: {data}")
        return {"error": "Processed request body is not a valid JSON object."}, 400

    student_name = data.get("student_name", "Student")
    title = data.get("title", "Untitled Project")
    content_raw = data.get("content", {}) 

    content = {}
    if isinstance(content_raw, str):
        try:
            content = json.loads(content_raw)
        except json.JSONDecodeError as e:
            app.logger.error(f"Error decoding content JSON: {e}")
            app.logger.error(f"Content received: {content_raw}")
            return {"error": "Invalid JSON format for 'content' field"}, 400
    else:
        content = content_raw

    if not isinstance(content, dict):
        app.logger.error(f"'content' field is not a dictionary: {type(content)}. Value: {content}")
        return {"error": "'content' field must be a dictionary"}, 400

    doc = Document()

    p1 = doc.add_paragraph()
    r1 = p1.add_run(f"Student name: {student_name}")
    r1.bold = True
    r1.font.size = Pt(20)

    p2 = doc.add_paragraph()
    r2 = p2.add_run(f"Title: {title}")
    r2.bold = True
    r2.font.size = Pt(20)

    doc.add_paragraph()

    for section, body in content.items():
        heading = doc.add_paragraph()
        heading_run = heading.add_run(f"{section}:")
        heading_run.bold = True
        heading_run.font.size = Pt(16)

        if isinstance(body, str):
            doc.add_paragraph(body)
        elif isinstance(body, dict):
            combined_text = ""
            if "text" in body and isinstance(body["text"], str):
                combined_text += body["text"]
            
            if combined_text.strip():
                doc.add_paragraph(combined_text.strip())

            if "bullets" in body and isinstance(body["bullets"], list):
                for bullet in body["bullets"]:
                    if isinstance(bullet, str):
                        bullet_paragraph = doc.add_paragraph(style='List Bullet')
                        bullet_paragraph.paragraph_format.left_indent = Inches(0.5)
                        bullet_paragraph.add_run(bullet)

    filename = f"{title.replace(' ', '_')}.docx"
    filepath = os.path.join("/tmp", filename)
    doc.save(filepath)

    # --- CRITICAL CHANGE HERE ---
    return send_file(
        filepath,
        as_attachment=True,
        download_name=filename, # Ensure the download name is set
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document' # Explicitly set mimetype
    )

@app.route("/", methods=["GET"])
def health():
    return "Docx Generator is running!"

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port)